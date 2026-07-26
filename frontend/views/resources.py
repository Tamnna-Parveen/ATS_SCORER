import io
import streamlit as st
import streamlit.components.v1 as components
from docx import Document
from docx.shared import Inches, Pt, RGBColor


# ---------------------------------------------------------
# HTML TEMPLATE GENERATORS (Full Resume Content)
# ---------------------------------------------------------

def get_template_1_sidebar() -> str:
    """Template 1: Left Dark Sidebar Two-Column Layout"""
    return """
    <!DOCTYPE html>
    <html>
    <head>
    <style>
        body { font-family: 'Georgia', serif; margin: 0; padding: 0; display: flex; background: #fff; color: #333; }
        .sidebar { width: 32%; background: #34495e; color: #fff; padding: 20px; box-sizing: border-box; min-height: 950px; }
        .main { width: 68%; padding: 25px; box-sizing: border-box; }
        .avatar { width: 85px; height: 85px; background: #bdc3c7; border-radius: 50%; margin-bottom: 15px; }
        .sidebar h3 { border-bottom: 1px solid #7f8c8d; padding-bottom: 5px; font-size: 14px; margin-top: 20px; color: #ecf0f1; text-transform: uppercase; }
        .sidebar p, .sidebar li { font-size: 11px; line-height: 1.6; word-break: break-word; }
        .sidebar ul { padding-left: 15px; margin: 0; }
        .name { font-size: 28px; color: #2c3e50; font-weight: bold; margin-bottom: 15px; }
        .section-title { font-size: 15px; color: #2c3e50; border-bottom: 1.5px solid #2c3e50; font-weight: bold; margin-top: 18px; padding-bottom: 2px; text-transform: uppercase; }
        .item-head { display: flex; justify-content: space-between; margin-top: 8px; font-size: 12px; }
        ul.main-list { font-size: 11.5px; padding-left: 15px; line-height: 1.5; margin-top: 5px; }
        li { margin-bottom: 4px; }
    </style>
    </head>
    <body>
        <div class="sidebar">
            <div class="avatar"></div>
            <h3>Contact</h3>
            <p>📞 +91 8758166872<br>✉️ emily.brown@sitare.org<br>🔗 linkedin.com/in/emilybrown<br>💻 github.com/emilybrown</p>
            
            <h3>Skills</h3>
            <ul>
                <li>Python</li><li>Java</li><li>JavaScript</li><li>HTML</li><li>CSS</li>
                <li>ReactJS</li><li>Flask</li><li>Django</li><li>MongoDB</li><li>MySQL</li>
                <li>PostgreSQL</li><li>Linux</li><li>NumPy</li><li>Pandas</li><li>Matplotlib</li>
                <li>Scikit-learn</li><li>NLP</li><li>Git</li><li>GitHub</li><li>Canva</li>
            </ul>
        </div>
        <div class="main">
            <div class="name">Emily Brown</div>
            
            <div class="section-title">Education</div>
            <div class="item-head"><b>Sitare University</b><i>May 2026</i></div>
            <div style="font-size: 11px; color: #555;">B.Tech, Computer Science, CGPA: 8.36 | Lucknow</div>
            <p style="font-size: 11px; margin-top: 4px;"><b>Relevant Coursework:</b> Machine Learning, Java Programming, Advanced Data Structures and Algorithms, Python Programming, Object Oriented Programming Concepts, Artificial Intelligence, Database Management Systems. <i>Honors: Dean's List</i></p>

            <div class="section-title">Experience</div>
            <div class="item-head"><b>FirstHive — Data Science Intern</b><i>May 2024 - Aug 2024</i></div>
            <div style="font-size: 11px; color: #555; font-style: italic;">Bangalore</div>
            <ul class="main-list">
                <li><b>Churn Prediction:</b> Challenges in customer retention due to lack of predictive insights and real-time analytics of customers.</li>
                <li><b>Model and Dashboard:</b> Built a customer churn prediction model (92% accuracy) and an interactive analytics dashboard.</li>
                <li><b>Better Data Insight:</b> Predicted churn, improving decision-making through real-time customer behavior tracking.</li>
                <li><b>Tools & Technologies Used:</b> Python, Scikit-learn, Pandas, Faker, Mimesis, Django, JavaScript, Chart.js, HTML, CSS.</li>
            </ul>

            <div class="section-title">Projects</div>
            <ul class="main-list">
                <li><b>SU_ChatBot:</b> Developed an AI-powered chatbot leveraging Llama API and pgVector to enable precise answers for questions asked. Used Flask, PostgreSQL (pgVector), Google Authentication, concurrent request handling and chat history tracking.</li>
                <li><b>HeartCheckAuto:</b> Developed a healthcare platform integrating an ML heart attack prediction model and with consultation features. Used ML, Google OAuth authentication, role-based access control, Flask, Chart.js and appointment scheduling.</li>
                <li><b>Library Management Portal:</b> Designed a Library Management Portal, enabling efficient book tracking and user records with real-time updates. Used tech stack: NoSQL (MongoDB) CRUD operations, Streamlit for interactive UI, and indexing for faster queries.</li>
            </ul>
        </div>
    </body>
    </html>
    """

def get_template_2_pill_blue() -> str:
    """Template 2: Top Contact Bar + Pill Badge Skills"""
    return """
    <!DOCTYPE html>
    <html>
    <head>
    <style>
        body { font-family: 'Segoe UI', Arial, sans-serif; padding: 20px; color: #2b2b2b; line-height: 1.4; }
        .name { font-size: 26px; font-weight: bold; color: #1e56a0; margin-bottom: 12px; }
        .contact-bar { background: #ebf3fa; padding: 8px; font-size: 12px; text-align: center; color: #163172; border-radius: 4px; }
        .section-title { font-size: 15px; font-weight: bold; color: #1e56a0; border-bottom: 2px solid #1e56a0; margin-top: 18px; padding-bottom: 3px; }
        .pills { display: flex; flex-wrap: wrap; gap: 6px; margin-top: 8px; }
        .pill { border: 1px solid #1e56a0; padding: 3px 9px; border-radius: 6px; font-size: 11px; color: #163172; background: #f6f9fc; font-weight: 500; }
        .badge { background: #1e56a0; color: #fff; padding: 2px 8px; border-radius: 10px; font-size: 10px; float: right; }
        ul { font-size: 12px; line-height: 1.5; padding-left: 18px; margin-top: 6px; }
        li { margin-bottom: 5px; }
    </style>
    </head>
    <body>
        <div class="name">Emily Brown</div>
        <div class="contact-bar">📞 +91 8758166872 &nbsp;|&nbsp; ✉️ emily.brown@sitare.org &nbsp;|&nbsp; 🔗 linkedin.com/in/emilybrown &nbsp;|&nbsp; 💻 GitHub</div>
        
        <div class="section-title">Skills</div>
        <div class="pills">
            <span class="pill">Python</span><span class="pill">Java</span><span class="pill">JavaScript</span>
            <span class="pill">HTML</span><span class="pill">CSS</span><span class="pill">ReactJS</span>
            <span class="pill">Flask</span><span class="pill">Django</span><span class="pill">MongoDB</span>
            <span class="pill">MySQL</span><span class="pill">PostgreSQL</span><span class="pill">Linux</span>
            <span class="pill">NumPy</span><span class="pill">Pandas</span><span class="pill">Matplotlib</span>
            <span class="pill">Scikit-learn</span><span class="pill">NLP</span><span class="pill">Git</span>
            <span class="pill">GitHub</span><span class="pill">Canva</span>
        </div>

        <div class="section-title">Education</div>
        <div style="margin-top:8px;">
            <span class="badge">May 2026</span>
            <b>Sitare University</b> — B.Tech, Computer Science, CGPA: 8.36
            <div style="font-size: 11px; color: #1e56a0; font-style: italic;">Lucknow</div>
        </div>
        <ul>
            <li><b>Relevant Coursework:</b> Machine Learning, Java Programming, Advanced Data Structures and Algorithms, Python Programming, Object Oriented Programming Concepts, Artificial Intelligence, Database Management Systems. <i>Honors: Dean's List</i></li>
        </ul>

        <div class="section-title">Experience</div>
        <div style="margin-top:8px;">
            <span class="badge">May 2024 - Aug 2024</span>
            <b>FirstHive</b> — Data Science Intern
            <div style="font-size: 11px; color: #1e56a0; font-style: italic;">Bangalore</div>
        </div>
        <ul>
            <li><b>Churn Prediction:</b> Challenges in customer retention due to lack of predictive insights and real-time analytics of customers.</li>
            <li><b>Model and Dashboard:</b> Built a customer churn prediction model (92% accuracy) and an interactive analytics dashboard.</li>
            <li><b>Better Data Insight:</b> Predicted churn, improving decision-making through real-time customer behavior tracking.</li>
            <li><b>Tools & Technologies Used:</b> Python, Scikit-learn, Pandas, Faker, Mimesis, Django, JavaScript, Chart.js, HTML, CSS.</li>
        </ul>

        <div class="section-title">Projects</div>
        <ul>
            <li><b>SU_ChatBot:</b> Developed an AI-powered chatbot leveraging Llama API and pgVector to enable precise answers for questions asked. Used Flask, PostgreSQL (pgVector), Google Authentication, concurrent request handling and chat history tracking.</li>
            <li><b>HeartCheckAuto:</b> Developed a healthcare platform integrating an ML heart attack prediction model and with consultation features. Used ML, Google OAuth authentication, role-based access control, Flask, Chart.js and appointment scheduling.</li>
            <li><b>Library Management Portal:</b> Designed a Library Management Portal, enabling efficient book tracking and user records with real-time updates. Used tech stack: NoSQL (MongoDB) CRUD operations, Streamlit for interactive UI, and indexing for faster queries.</li>
        </ul>
    </body>
    </html>
    """

def get_template_3_grey_pills() -> str:
    """Template 3: Grey Filled Soft Pills Layout"""
    return """
    <!DOCTYPE html>
    <html>
    <head>
    <style>
        body { font-family: 'Helvetica Neue', Arial, sans-serif; padding: 20px; color: #222; }
        .name { font-size: 28px; font-weight: 800; color: #111; }
        .contact-info { background: #f1f5f9; padding: 8px 12px; font-size: 12px; border-radius: 6px; margin-top: 10px; }
        .section-title { font-size: 14px; font-weight: 800; letter-spacing: 1px; color: #0f172a; margin-top: 20px; border-bottom: 1px solid #cbd5e1; padding-bottom: 4px; text-transform: uppercase; }
        .pills { display: flex; flex-wrap: wrap; gap: 6px; margin-top: 8px; }
        .pill { background: #e2e8f0; color: #334155; padding: 3px 10px; border-radius: 12px; font-size: 11px; font-weight: 600; }
        ul { font-size: 12px; line-height: 1.5; padding-left: 16px; margin-top: 6px; }
        li { margin-bottom: 5px; }
    </style>
    </head>
    <body>
        <div class="name">Emily Brown</div>
        <div class="contact-info">✉️ emily.brown@sitare.org | 📞 +91 8758166872 | 🔗 linkedin.com/in/emilybrown | 💻 GitHub</div>
        
        <div class="section-title">SKILLS</div>
        <div class="pills">
            <span class="pill">Python</span><span class="pill">Java</span><span class="pill">JavaScript</span>
            <span class="pill">HTML</span><span class="pill">CSS</span><span class="pill">ReactJS</span>
            <span class="pill">Flask</span><span class="pill">Django</span><span class="pill">MongoDB</span>
            <span class="pill">MySQL</span><span class="pill">PostgreSQL</span><span class="pill">Linux</span>
            <span class="pill">NumPy</span><span class="pill">Pandas</span><span class="pill">Matplotlib</span>
            <span class="pill">Scikit-learn</span><span class="pill">NLP</span><span class="pill">Git</span>
            <span class="pill">GitHub</span><span class="pill">Canva</span>
        </div>

        <div class="section-title">EDUCATION</div>
        <div style="font-size:13px; font-weight:bold; margin-top:8px;">
            B.Tech, Computer Science, CGPA: 8.36 <span style="float:right; font-weight:normal; color:#666;">May 2026</span>
        </div>
        <div style="font-size:11px; color:#555; font-style:italic;">Sitare University, Lucknow</div>
        <ul>
            <li>Relevant Coursework: Machine Learning, Java Programming, Advanced Data Structures and Algorithms, Python Programming, Object Oriented Programming Concepts, Artificial Intelligence, Database Management Systems. Honors: Dean's List</li>
        </ul>

        <div class="section-title">EXPERIENCE</div>
        <div style="font-size:13px; font-weight:bold; margin-top:8px;">
            Data Science Intern <span style="float:right; font-weight:normal; color:#666;">May 2024 - Aug 2024</span>
        </div>
        <div style="font-size:11px; color:#555; font-style:italic;">FirstHive, Bangalore</div>
        <ul>
            <li>Churn Prediction: Challenges in customer retention due to lack of predictive insights and real-time analytics of customers.</li>
            <li>Model and Dashboard: Built a customer churn prediction model (92% accuracy) and an interactive analytics dashboard.</li>
            <li>Better Data Insight: Predicted churn, improving decision-making through real-time customer behavior tracking.</li>
            <li>Tools Used: Python, Scikit-learn, Pandas, Faker, Mimesis, Django, JavaScript, Chart.js, HTML, CSS.</li>
        </ul>

        <div class="section-title">PROJECTS</div>
        <ul>
            <li><b>SU_ChatBot:</b> Developed an AI-powered chatbot leveraging Llama API and pgVector to enable precise answers for questions asked. Used Flask, PostgreSQL (pgVector), Google Authentication, concurrent request handling and chat history tracking.</li>
            <li><b>HeartCheckAuto:</b> Developed a healthcare platform integrating an ML heart attack prediction model and with consultation features. Used ML, Google OAuth authentication, role-based access control, Flask, Chart.js and appointment scheduling.</li>
            <li><b>Library Management Portal:</b> Designed a Library Management Portal, enabling efficient book tracking and user records with real-time updates. Used tech stack: NoSQL (MongoDB) CRUD operations, Streamlit for interactive UI, and indexing for faster queries.</li>
        </ul>
    </body>
    </html>
    """

def get_template_4_serif_centered() -> str:
    """Template 4: Academic Serif Style"""
    return """
    <!DOCTYPE html>
    <html>
    <head>
    <style>
        body { font-family: 'Times New Roman', Times, serif; padding: 20px; color: #111; line-height: 1.4; }
        .name { font-size: 28px; text-align: left; font-weight: bold; }
        .contact { font-size: 12px; margin-top: 5px; }
        .section-head { background: #e8f1f5; color: #1a365d; text-align: center; font-weight: bold; font-size: 13px; letter-spacing: 1px; padding: 4px; margin-top: 15px; text-transform: uppercase; }
        ul { font-size: 12px; line-height: 1.5; padding-left: 18px; margin-top: 6px; }
        li { margin-bottom: 5px; }
    </style>
    </head>
    <body>
        <div class="name">Emily Brown</div>
        <div class="contact">✉️ emily.brown@sitare.org | 📞 +91 8758166872 | 🔗 linkedin.com/in/emilybrown | 💻 GitHub</div>
        
        <div class="section-head">SKILLS</div>
        <p style="font-size:12px; margin: 8px 0; text-align:center;">
            Python • Java • JavaScript • HTML • CSS • ReactJS • Flask • Django • MongoDB • MySQL • PostgreSQL • Linux • NumPy • Pandas • Matplotlib • Scikit-learn • NLP • Git • GitHub • Canva
        </p>

        <div class="section-head">EDUCATION</div>
        <div style="margin-top:8px; font-size:12px;">
            <b>Sitare University</b> <span style="float:right; font-style:italic;">Lucknow, May 2026</span><br>
            <i>B.Tech, Computer Science, CGPA: 8.36</i>
        </div>
        <ul>
            <li>Relevant Coursework: Machine Learning, Java Programming, Advanced Data Structures and Algorithms, Python Programming, Object Oriented Programming Concepts, Artificial Intelligence, Database Management Systems. Honors: Dean's List</li>
        </ul>

        <div class="section-head">EXPERIENCE</div>
        <div style="margin-top:8px; font-size:12px;">
            <b>FirstHive</b> — <i>Data Science Intern</i> <span style="float:right; font-style:italic;">Bangalore, May 2024 - Aug 2024</span>
        </div>
        <ul>
            <li>Churn Prediction: Challenges in customer retention due to lack of predictive insights and real-time analytics of customers.</li>
            <li>Model and Dashboard: Built a customer churn prediction model (92% accuracy) and an interactive analytics dashboard.</li>
            <li>Better Data Insight: Predicted churn, improving decision-making through real-time customer behavior tracking.</li>
            <li>Tools & Technologies Used: Python, Scikit-learn, Pandas, Faker, Mimesis, Django, JavaScript, Chart.js, HTML, CSS.</li>
        </ul>

        <div class="section-head">PROJECTS</div>
        <ul>
            <li><b>SU_ChatBot:</b> Developed an AI-powered chatbot leveraging Llama API and pgVector to enable precise answers for questions asked. Used Flask, PostgreSQL (pgVector), Google Authentication, concurrent request handling and chat history tracking.</li>
            <li><b>HeartCheckAuto:</b> Developed a healthcare platform integrating an ML heart attack prediction model and with consultation features. Used ML, Google OAuth authentication, role-based access control, Flask, Chart.js and appointment scheduling.</li>
            <li><b>Library Management Portal:</b> Designed a Library Management Portal, enabling efficient book tracking and user records with real-time updates. Used tech stack: NoSQL (MongoDB) CRUD operations, Streamlit for interactive UI, and indexing for faster queries.</li>
        </ul>
    </body>
    </html>
    """

def get_template_5_blue_header_tag() -> str:
    """Template 5: Blue Tag Card Header Style"""
    return """
    <!DOCTYPE html>
    <html>
    <head>
    <style>
        body { font-family: 'Segoe UI', Tahoma, sans-serif; padding: 20px; color: #222; }
        .header-tag { background: #3b69b5; color: #fff; padding: 12px 24px; border-radius: 6px 18px 6px 6px; font-size: 26px; font-weight: bold; display: inline-block; }
        .contact-list { font-size: 11.5px; margin-top: 12px; color: #444; }
        .section-title { color: #3b69b5; font-size: 15px; font-weight: bold; border-bottom: 1.5px solid #3b69b5; margin-top: 18px; padding-bottom: 2px; }
        ul { font-size: 12px; line-height: 1.5; padding-left: 18px; margin-top: 6px; }
        li { margin-bottom: 5px; }
    </style>
    </head>
    <body>
        <div class="header-tag">Emily Brown</div>
        <div class="contact-list">
            📞 +91 8758166872 &nbsp;|&nbsp; ✉️ emily.brown@sitare.org &nbsp;|&nbsp; 🔗 linkedin.com/in/emilybrown &nbsp;|&nbsp; 💻 GitHub
        </div>

        <div class="section-title">🎓 Education</div>
        <div style="margin-top:6px; font-size:12px;">
            <b>Sitare University</b>, Lucknow <span style="float:right; color:#3b69b5; font-weight:600;">May 2026</span><br>
            <i>B.Tech, Computer Science, CGPA: 8.36</i>
        </div>
        <ul>
            <li><b>Relevant Coursework:</b> Machine Learning, Java Programming, Advanced Data Structures and Algorithms, Python Programming, Object Oriented Programming Concepts, Artificial Intelligence, Database Management Systems. <i>Honors: Dean's List</i></li>
        </ul>

        <div class="section-title">💼 Experience</div>
        <div style="margin-top:6px; font-size:12px;">
            <b>FirstHive</b>, Bangalore — <i>Data Science Intern</i> <span style="float:right; color:#3b69b5; font-weight:600;">May 2024 - Aug 2024</span>
        </div>
        <ul>
            <li><b>Churn Prediction:</b> Challenges in customer retention due to lack of predictive insights and real-time analytics of customers.</li>
            <li><b>Model and Dashboard:</b> Built a customer churn prediction model (92% accuracy) and an interactive analytics dashboard.</li>
            <li><b>Better Data Insight:</b> Predicted churn, improving decision-making through real-time customer behavior tracking.</li>
            <li><b>Tools Used:</b> Python, Scikit-learn, Pandas, Faker, Mimesis, Django, JavaScript, Chart.js, HTML, CSS.</li>
        </ul>

        <div class="section-title">🚀 Projects</div>
        <ul>
            <li><b>SU_ChatBot:</b> Developed an AI-powered chatbot leveraging Llama API and pgVector to enable precise answers for questions asked. Used Flask, PostgreSQL (pgVector), Google Authentication, concurrent request handling and chat history tracking.</li>
            <li><b>HeartCheckAuto:</b> Developed a healthcare platform integrating an ML heart attack prediction model and with consultation features. Used ML, Google OAuth authentication, role-based access control, Flask, Chart.js and appointment scheduling.</li>
            <li><b>Library Management Portal:</b> Designed a Library Management Portal, enabling efficient book tracking and user records with real-time updates. Used tech stack: NoSQL (MongoDB) CRUD operations, Streamlit for interactive UI, and indexing for faster queries.</li>
        </ul>
    </body>
    </html>
    """

# ---------------------------------------------------------
# DOCX GENERATOR FUNCTION
# ---------------------------------------------------------

def create_selected_docx(template_name: str) -> bytes:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    # Name
    p_name = doc.add_paragraph()
    r_name = p_name.add_run("Emily Brown")
    r_name.font.size = Pt(22)
    r_name.font.bold = True
    
    if "Blue" in template_name:
        r_name.font.color.rgb = RGBColor(30, 86, 160)
    elif "Sidebar" in template_name:
        r_name.font.color.rgb = RGBColor(52, 73, 94)

    # Contact Info
    p_contact = doc.add_paragraph()
    r_contact = p_contact.add_run("📞 +91 8758166872 | ✉️ emily.brown@sitare.org | 🔗 linkedin.com/in/emilybrown | 💻 GitHub")
    r_contact.font.size = Pt(9.5)

    def add_heading(title: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(title.upper())
        r.font.size = Pt(12)
        r.font.bold = True
        r.font.color.rgb = RGBColor(30, 86, 160)

    # Skills
    add_heading("Skills")
    p_skills = doc.add_paragraph()
    r_skills = p_skills.add_run("Python • Java • JavaScript • HTML • CSS • ReactJS • Flask • Django • MongoDB • MySQL • PostgreSQL • Linux • NumPy • Pandas • Matplotlib • Scikit-learn • NLP • Git • GitHub • Canva")
    r_skills.font.size = Pt(9.5)

    # Education
    add_heading("Education")
    p_edu = doc.add_paragraph()
    r_edu = p_edu.add_run("Sitare University — B.Tech, Computer Science (CGPA: 8.36) | Lucknow (May 2026)")
    r_edu.font.bold = True
    r_edu.font.size = Pt(10)

    p_edu_cw = doc.add_paragraph(style='List Bullet')
    r_edu_cw = p_edu_cw.add_run("Relevant Coursework: Machine Learning, Java Programming, Advanced Data Structures and Algorithms, Python Programming, Object Oriented Programming Concepts, Artificial Intelligence, Database Management Systems. Honors: Dean's List")
    r_edu_cw.font.size = Pt(9)

    # Experience
    add_heading("Experience")
    p_exp = doc.add_paragraph()
    r_exp = p_exp.add_run("FirstHive — Data Science Intern | Bangalore (May 2024 - Aug 2024)")
    r_exp.font.bold = True
    r_exp.font.size = Pt(10)

    exp_bullets = [
        "Churn Prediction: Challenges in customer retention due to lack of predictive insights and real-time analytics of customers.",
        "Model and Dashboard: Built a customer churn prediction model (92% accuracy) and an interactive analytics dashboard.",
        "Better Data Insight: Predicted churn, improving decision-making through real-time customer behavior tracking.",
        "Tools & Technologies Used: Python, Scikit-learn, Pandas, Faker, Mimesis, Django, JavaScript, Chart.js, HTML, CSS."
    ]
    for b in exp_bullets:
        p_b = doc.add_paragraph(style='List Bullet')
        r_b = p_b.add_run(b)
        r_b.font.size = Pt(9)

    # Projects
    add_heading("Projects")
    proj_bullets = [
        "SU_ChatBot: Developed an AI-powered chatbot leveraging Llama API and pgVector to enable precise answers for questions asked. Used Flask, PostgreSQL (pgVector), Google Authentication, concurrent request handling and chat history tracking.",
        "HeartCheckAuto: Developed a healthcare platform integrating an ML heart attack prediction model and with consultation features. Used ML, Google OAuth authentication, role-based access control, Flask, Chart.js and appointment scheduling.",
        "Library Management Portal: Designed a Library Management Portal, enabling efficient book tracking and user records with real-time updates. Used tech stack: NoSQL (MongoDB) CRUD operations, Streamlit for interactive UI, and indexing for faster queries."
    ]
    for pb in proj_bullets:
        p_pb = doc.add_paragraph(style='List Bullet')
        r_pb = p_pb.add_run(pb)
        r_pb.font.size = Pt(9)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ---------------------------------------------------------
# MAIN STREAMLIT RENDER
# ---------------------------------------------------------

def render() -> None:
    st.markdown("""
    <style>
        .hero-banner {
            background: linear-gradient(135deg, #1e3a8a 0%, #3b82f6 100%);
            padding: 24px;
            border-radius: 14px;
            color: #ffffff;
            margin-bottom: 24px;
        }
    </style>
    <div class="hero-banner">
        <h2 style="margin:0; color:#fff;">📑 Resume Template Library</h2>
        <p style="margin:4px 0 0 0; color:#e0e7ff; font-size:14px;">Select any format, preview live, and download instantly in .DOCX or .HTML.</p>
    </div>
    """, unsafe_allow_html=True)

    # Template Options
    templates_dict = {
        "1. Dark Sidebar Two-Column Layout": get_template_1_sidebar(),
        "2. Modern Blue Accent with Skill Pills": get_template_2_pill_blue(),
        "3. Minimalist Soft Grey Badges": get_template_3_grey_pills(),
        "4. Classic Serif Block Heading Layout": get_template_4_serif_centered(),
        "5. Blue Top-Left Tag Header Style": get_template_5_blue_header_tag()
    }

    selected_template_name = st.selectbox(
        "🎨 **Choose a Resume Template Format:**",
        list(templates_dict.keys())
    )

    selected_html = templates_dict[selected_template_name]
    docx_file = create_selected_docx(selected_template_name)

    st.markdown("---")

    col_preview, col_download = st.columns([2, 1], gap="large")

    with col_preview:
        st.subheader("👁️ Live Preview")
        # Live HTML Preview Rendering
        components.html(selected_html, height=720, scrolling=True)

    with col_download:
        st.subheader("📥 Download File")
        st.info(f"Selected Layout:\n**{selected_template_name.split('.')[1]}**")

        st.markdown("<br>", unsafe_allow_html=True)

        st.download_button(
            label="📄 Download Word (.DOCX)",
            data=docx_file,
            file_name=f"{selected_template_name.split(' ')[1]}_Resume_Template.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            type="primary"
        )

        st.markdown("<div style='height: 10px;'></div>", unsafe_allow_html=True)

        st.download_button(
            label="🌐 Download Web (.HTML)",
            data=selected_html,
            file_name=f"{selected_template_name.split(' ')[1]}_Resume_Template.html",
            mime="text/html",
            use_container_width=True
        )

        st.markdown("""
        ---
        **Quick Summary:**
        - ✅ 100% ATS-Friendly Structure
        - ✏️ Fully Editable Document
        - 📱 Instant Interactive Preview
        """)

    # ---------------------------------------------------------
    # SPACIOUS FULL-WIDTH DO'S AND DON'TS SECTION (PLACED BELOW)
    # ---------------------------------------------------------
    st.markdown("<br><hr>", unsafe_allow_html=True)
    st.subheader("💡 Best Practices & Formatting Rules")
    
    col_dos, col_donts = st.columns(2, gap="large")
    
    with col_dos:
        st.success("### ✅ DO's")
        st.markdown("""
        * **Standard Fonts:** Use clean, readable fonts (e.g., Arial, Georgia, Segoe UI).
        * **Action Verbs & Impact:** Begin bullet points with strong verbs (e.g., *Developed*, *Built*, *Optimized*).
        * **Measurable Metrics:** Quantify achievements wherever possible (e.g., *92% accuracy*, *CGPA 8.36*).
        * **Clear Hierarchy:** Ensure bold section titles and distinct date alignments.
        * **Consistent Spacing:** Keep margins and bullet paddings uniform throughout.
        """)
        
    with col_donts:
        st.error("### ❌ DON'Ts")
        st.markdown("""
        * **No Images for Text:** Avoid putting critical text inside graphics or unparseable text boxes.
        * **Avoid Over-Styling:** Don't use too many colors or overly decorative font styles.
        * **No Missing Dates:** Don't leave dates or locations unaligned/missing.
        * **Avoid Spelling Errors:** Thoroughly proofread tech stacks and project descriptions.
        * **No Walls of Text:** Keep descriptions structured as bullet points instead of long paragraphs.
        """)